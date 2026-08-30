# -*- coding: utf-8 -*-
"""
明鉴 · 创新竞赛评审打分 API（v1.8.0 引擎）
==========================================
把「明鉴」评审专家（innovation-review-expert v1.8.0）封装为 HTTP 打分服务：

    POST /api/review   上传赛道 + 计划书文本 → 返回结构化评分 JSON
    GET  /api/tracks   赛道列表（前端下拉框）
    GET  /api/health   健康检查

DeepSeek API Key 只存在于本服务环境变量中，绝不进入前端代码。
部署见同目录 README.md。
"""

import os
import re
import io
import json
import asyncio
import logging
from datetime import date
from urllib.parse import quote

import httpx
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel, Field

# python-docx（Word 报告导出）
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None


def _to_rgb(color):
    """兼容 int(0xRRGGBB) 与 RGBColor 两种传参"""
    if isinstance(color, int):
        return RGBColor((color >> 16) & 0xFF, (color >> 8) & 0xFF, color & 0xFF)
    return color

# ─────────────────────────── 配置 ───────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_PATH = os.path.join(BASE_DIR, "reviewer_data.json")

DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "").strip()
DEEPSEEK_URL = os.environ.get("DEEPSEEK_URL", "https://api.deepseek.com/chat/completions").strip()
DEEPSEEK_MODEL = os.environ.get("DEEPSEEK_MODEL", "deepseek-chat").strip()
MAX_INPUT_CHARS = int(os.environ.get("MAX_INPUT_CHARS", "95000"))   # 中文字符预算（DeepSeek 64K 上下文：中文约0.6 token/字，9.5万字符≈5.7万 token + 输出4K ≈ 6.1万，实测安全）
TIMEOUT = float(os.environ.get("API_TIMEOUT", "120"))

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger("mingjian")

# 10 个赛道 key，顺序与前端 TRACK_KEYS 一致
TRACK_KEYS = [
    "gaojiao_chuangyi", "gaojiao_chuangye", "honglv_gongyi", "honglv_chuangyi",
    "honglv_chuangye", "zhijiao_chuangyi", "zhijiao_chuangye",
    "chanye_qimingti", "chanye_chengguo", "mengya",
]

# ─────────────────────────── 数据加载 ───────────────────────────
with open(DATA_PATH, encoding="utf-8") as f:
    REVIEWER_DATA = json.load(f)

# ─────────────────── 明鉴 v1.8.0 系统提示词（V3.1 五步打分法） ───────────────────
SYSTEM_PROMPT = """你是「明鉴」，一位深耕创新领域二十年的资深评审专家，长期担任各类创新竞赛评审委员，阅遍成千上万份创新作品。你已内化《中国国际大学生创新竞赛》全部官方评审规则（10 个赛道/组别），并拥有 155+ 个真实获奖案例（2019-2025，国金/国银/国铜/省奖全档位）校准出的实证锚点库。

【本次任务】用户会给你一份商业计划书文本，以及指定赛道的官方评分规则（含维度、权重、逐级评分标准）和该赛道的获奖案例锚点库。你严格按这些材料评审，输出结构化 JSON 评分结果。

【评审方法论 · V3.1 五步打分法（条例定刻度，案例定分寸）】
① 案例定分寸（强制先行，最高纪律）：评分前必须先精读 user prompt 中的【实证校准 · 档位定位器与临界规律】与【获奖案例锚点库】——"这个档位的证据长什么样"（三件套/四件套达标到什么程度算国金、国银、国铜），按档位快速定位器五条线（公司状态/验证层级/落地金额/营收呈现/量化指标）先定位，再逐维度打分。条例定刻度 × 案例定分寸缺一不可，禁止只按规则刻度硬套（实证教训：e-spider 首评未查画像偏严 9.5 分，实际国银）。
② 核心链定档：核心证据链达标情况直接定档位。创意组"三件套"（天花板证据/权威验证/真实落地）、创业组"经营四件套"（公司/到账营收/融资/实名客户）、红旅"公益五维"。核心链达标即进档位基线，不再因单项短板逐项扣分——金奖不可能条条规则满足。
③ 强项定位：核心强项（验证硬度/专家层级/落地金额/团队稀缺度）达"突出标准"→ 档内上移/加分溢出。
④ 短板豁免/致命降档：短板命中"允许短板清单"不扣分（创意组：无营收/无订单、未注册公司、专利少或零、育人叙事弱、附件缺失；创业组：营收百万级/未盈利、融资仅天使轮、育人叙事弱、专利归校、非顶尖高校、客户未实名但有具名表述；红旅公益：无经营实绩、覆盖面小但有深度案例）；只有命中"致命短板清单"才降档（纯概念无验证/无样机、四件套缺两件以上、材料自相矛盾致可信度崩溃、技术路线无原创性证据、团队与项目完全脱节）。
⑤ 材料完整度小扣（独立于内容分）：缺核心章节扣 1-3 分；正文乱码/数据矛盾扣 0.5-1 分；流传渠道特性（水印/纯图片可 OCR/附件缺失）不扣分。
【执行纪律（四必须）】
1. 评分前必查实证校准段与锚点画像（先"定分寸"再"定刻度"细节）；
2. 档位是竞争性结果，不可机械判档——存在"强数据落低档"（神州VR 营收 2400 万+订单 6000 万仍国铜、授虾以蚕 1.5 亿营业额仍国银），赛道天花板、材料可信度是隐性因素，须核查后说明；
3. 每条失分必须配可落地的改进建议（具体到章节/页面：改哪里、怎么改、改完能提升多少），禁止只判档不给方；
4. 评分尺度放宽（V2.6 实证沉淀）：附件原件缺失降为提示项不扣分（表述具体+数字明确+机构具名即按已呈现证据计分）；纯图片 OCR 可完整提取不扣分；水印不扣分；预测性财务无测算不扣分（仅提示）；子项未提及按满分 10-15% 给基线分；档间取值取证据上限。禁止反向贴分（材料中不存在的证据不得凭空加分）。

【三档分界速记】国铜=无实物/纯计划/验证无动作；国银=实物+试用+意向/协议动作齐全（未授权/无到账/无金额不降档）；国金=授权/查新/到账顶格；省金=晋级标签按材料质量独立打分（达国奖水平给国奖区间，不因标签压分）。

【八大铁律 · 必须遵守】
1. 硬证据 > 叙事："世界首创/领航者/纪录"等标题不计分；注册证/审定证书/检测报告/到账营收/合同/获奖才计分。
2. 证据错位识别：论文/专利/学术实验是"学术证据"，撑底线分不撑项目分；项目级验证 = 产品原型/临床/试用/检测/订单/合同。
3. 只认"到账营收"，不认"预计/订单/科研经费"；以"户/家/单位"计比"人次"更有说服力。
4. 态度分陷阱：选题方向、财务克制、学校平台不构成加分。
5. 赛制层级：省金=省级一等奖+国赛入围资格，与国奖是晋级关系而非并列档位，按材料质量打分不因标签压分。
6. 格式合规酌情扣分：缺件/缺页/水印残留/纯图片无文字/排版混乱/封面信息缺失，在材料呈现维度酌情扣分（不构成降档，累计不超该维度满分 20-30%；缺核心章节致证据无法核验时可影响档位）。
7. 阶段适配豁免：仅长周期硬科技早期项目（芯片/新药等）在满足"天花板团队+架构级创新+战略窗口"时可豁免当期经营实绩，消费级/软件/服务类一律不豁免。
8. 严格区分赛道：不同赛道维度权重不同，必须用当前指定赛道的规则，禁止混用。

【七档标尺 V3.1（实证版）】
国金 85-100（三件套/四件套齐备且顶格：查新首创/顶会/院士源头 + 国家级验证 + 大额到账订单 + 实名客户）/ 国银 76-85（证据链完整，缺"行业第一"级证据）/ 国铜 68-76（拟成立/未注册公司、验证无动作、计划无数额）/ 省金 64-68（晋级标签非档次标签：按材料质量独立打分，质量上限可至国金）/ 省银 58-64（有公司/标准/产业依托，但项目自身硬证据不足）/ 省铜 52-58（未注册公司/小体量/预测收入）/ 未获奖 35-52（零验证/纯方案/凑分书，低于 52 一律不算获奖）。

【双轨评分】total 为材料分（评审对象=所提交材料，按实际呈现证据打分）。可额外给出 work_score 作品水平分（评审对象=作品，依据可见证据上限+锚点实证记录推演），并在 summary 或 potential 中说明两分差 = 材料版本落差 + 答辩补充；作品水平分禁止"因国金标签贴分"。

【材料类型】默认按 BP（商业计划书）口径评审。用户提供的是纯文本（可能因 OCR 或粘贴丢失排版），不因文本格式简陋而惩罚内容，但可指出"证据无法核验"的缺失项。

【输出要求 · 严格 JSON】
只输出一个合法 JSON 对象，不要输出任何解释、markdown 代码块标记、或 JSON 之外的文字。字段 schema 如下：
{
  "total": 数字,               // 总分（材料分），保留1位小数
  "work_score": 数字,          // 可选：作品水平分（保留1位小数）
  "level": 字符串,             // 档位名：国金/国银/国铜/省金/省银/省铜/未获奖
  "level_range": 字符串,       // 如 "68-76"
  "summary": 字符串,           // 一句话定调（犀利、有记忆点）
  "max_leverage": 字符串,      // 最大改造杠杆（最高性价比改进方向）
  "dims": [                    // 每个一级维度一条
    {
      "name": "个人成长", "full": 25, "score": 20.0, "rate": 80,
      "diagnosis": "一句话诊断（含证据）",
      "subdims": [{"name":"立德树人","full":4,"score":3.5,"gain":"得分点：材料中的证据（章节/数据/证书）","loss":"扣分点：缺什么/扣在哪"}]
    }
  ],
  "highlights": ["三大亮点之一，具体到数据/证书/章节", "…", "…"],
  "weaknesses": ["三大不足之一，说明约失多少分、涉及哪些维度", "…", "…"],
  "suggestions": [{"priority":1,"text":"最高优先级改进建议，具体到章节/页面可落地"}, {"priority":2,"text":"…"}],
  "potential": "获奖潜力判断：当前定位 + 改造后潜力"
}

【评分要求】
- 逐维度、逐子维度打分，每项得分必须能从规则层级 + 作品证据中找到依据。
- 子维度得分之和应等于该维度得分；各维度得分之和等于 total。
- **每个子维度必须同时给出「得分点 gain」（材料中支持该得分的证据，如章节/页码/数据/证书）与「扣分点 loss」（缺什么/扣在哪）**；无失分项 loss 写"无"；材料未呈现的证据不得编造写入 gain。
- **total 必须落在 level 对应档位的 level_range 区间内**（如档位=省金则 total ∈ [64,68]）；若按证据计算的总分与档位区间有出入，以档位为准微调 total 并同步调整相关维度得分，保证三者自洽。
- 若材料信息不足，在 diagnosis / weaknesses 中明确列出缺失项，不得臆测补足证据。
- 批评必须附改进路径，杜绝空泛的"需加强"式表述。"""


# ─────────────────────────── prompt 组装（移植前端 buildAiPrompt） ───────────────────────────

# 超长文档分段评审：每段长度与信息提取提示词
SEGMENT_CHARS = 15000  # 每段约 1.5 万字符（单段可完整阅读）
EXTRACT_PROMPT = """你是商业计划书信息提取器。请阅读下方计划书片段，提取其中实际出现的信息，只输出一个 JSON 对象，字段如下：
{"项目名称":"", "项目背景与痛点":"", "产品与服务":"", "技术创新与壁垒":"", "市场分析":"", "商业模式与收入":"", "财务与融资":"", "团队与落地证据":"", "其他关键信息":""}
要求：1) 只提取该片段中真实出现的内容，保留具体数据（金额/数量/日期/百分比）；2) 某字段片段中未提及则写"未提及"；3) 不要总结观点，忠实摘录事实。"""


def body_budget_for(track_key: str) -> int:
    """计算某赛道留给计划书正文的字符预算"""
    track = REVIEWER_DATA["tracks"].get(track_key) or {}
    global_data = REVIEWER_DATA.get("global", {})
    fixed_len = (
        len(SYSTEM_PROMPT)
        + len(global_data.get("manual_v31", ""))
        + len(global_data.get("anchor_index", ""))
        + len(global_data.get("bp_ppt_rules", ""))
        + 1000
    )
    budget = MAX_INPUT_CHARS - fixed_len - len(track.get("rule", ""))
    a_max = min(len(track.get("anchors", "")), max(6000, int(budget * 0.42)))
    return max(8000, budget - a_max)


async def _llm_once(client: httpx.AsyncClient, messages: list, max_tokens: int = 800):
    """调用 DeepSeek 并返回文本内容"""
    resp = await client.post(
        DEEPSEEK_URL,
        headers={
            "Content-Type": "application/json",
            "Authorization": "Bearer " + DEEPSEEK_API_KEY,
        },
        json={
            "model": DEEPSEEK_MODEL,
            "messages": messages,
            "temperature": 0.1,
            "max_tokens": max_tokens,
            "response_format": {"type": "json_object"},
            "stream": False,
        },
    )
    if resp.status_code != 200:
        raise HTTPException(status_code=502, detail=f"上游模型返回 HTTP {resp.status_code}")
    content = resp.json().get("choices", [{}])[0].get("message", {}).get("content", "")
    if not content:
        raise HTTPException(status_code=502, detail="模型返回内容为空")
    return content


async def _extract_segment(client: httpx.AsyncClient, seg: str):
    """提取单个片段的要点 JSON"""
    try:
        content = await _llm_once(
            client,
            [
                {"role": "system", "content": EXTRACT_PROMPT},
                {"role": "user", "content": seg},
            ],
            max_tokens=1000,
        )
        return extract_json(content)
    except Exception:
        return {}  # 单段失败不阻塞整体，缺失信息在要点包中留空


async def segment_review_text(text: str):
    """超长文档：切段→并发提取→汇总要点包。返回 (要点包文本, 段数)"""
    chunks = [text[i : i + SEGMENT_CHARS] for i in range(0, len(text), SEGMENT_CHARS)]
    async with httpx.AsyncClient(timeout=TIMEOUT) as client:
        results = await asyncio.gather(*[_extract_segment(client, c) for c in chunks])

    merged: dict = {}
    for r in results:
        if not isinstance(r, dict):
            continue
        for k, v in r.items():
            if isinstance(v, str) and v and v != "未提及":
                merged[k] = (merged.get(k, "") + "；" + v).strip("；")
    merged = {k: v for k, v in merged.items() if v}

    if not merged:
        # 全段提取失败，退化为保头保尾原文
        return text, len(chunks)

    lines = [f"【{k}】{v}" for k, v in merged.items()]
    summary = (
        "以下为超长计划书的分段提取要点（原文共 %d 段，已逐段精读并汇总），"
        "请基于全部要点做完整评审，注意信息完整性：\n\n" % len(chunks)
    ) + "\n".join(lines)
    return summary, len(chunks)


def build_prompt(track_key: str, text: str) -> str:
    track = REVIEWER_DATA["tracks"].get(track_key)
    if not track:
        raise HTTPException(status_code=400, detail=f"未知赛道: {track_key}")

    global_data = REVIEWER_DATA.get("global", {})
    sys_len = len(SYSTEM_PROMPT)
    manual = global_data.get("manual_v31", "")       # 实证校准：档位定位器+临界规律+放宽规则
    index = global_data.get("anchor_index", "")
    bprules = global_data.get("bp_ppt_rules", "")
    fixed_len = sys_len + len(manual) + len(index) + len(bprules) + 1000  # 分隔/标题余量

    budget = MAX_INPUT_CHARS - fixed_len

    # 1) 官方评分规则：优先完整保留（打分依据）
    rule = track.get("rule", "")
    budget -= len(rule)
    if budget < 6000:
        rule = rule[: max(6000, budget + len(rule))]
        budget = max(0, budget)

    # 2) 锚点库：截取前部（国金档在前），至少保留 6000 字；正文优先（锚点取预算 42%，全量优先）
    anchors = track.get("anchors", "")
    a_max = min(len(anchors), max(6000, int(budget * 0.42)))
    anchors = anchors[:a_max]
    budget -= a_max

    # 3) 计划书正文：保头保尾智能截断（财务/团队/落地计划通常在文档尾部，不能丢）
    body_limit = max(8000, budget)
    truncated = False
    if len(text) > body_limit:
        truncated = True
        head = int(body_limit * 0.62)
        tail = body_limit - head - 60
        if tail < 1200:
            head = body_limit - 1200
            tail = 1200
        cut_n = len(text) - (head + tail)
        body = (
            text[:head]
            + f"\n\n……（原文共 {len(text)} 字，因长度限制省略中间约 {cut_n} 字，以下为文末部分，请一并评审）……\n\n"
            + text[-tail:]
        )
    else:
        body = text

    prompt = "\n".join([
        "【指定赛道】" + track.get("name", track_key),
        "",
        "【官方评分规则（条例定刻度）】",
        rule,
        "",
        "【实证校准 · 档位定位器与临界规律（评分前必读，案例定分寸）】",
        manual,
        "",
        "【获奖案例锚点库（案例参照）】",
        anchors,
        "",
        "【七档标尺与通用通则】",
        index,
        "",
        "【材料呈现规则】",
        bprules,
        "",
        "【待评审的商业计划书正文】",
        body,
        "",
        "请严格按上述规则与锚点库评审，只输出一个合法 JSON 对象，不要输出任何其他文字。",
    ])
    return prompt, {"truncated": truncated, "input_chars": len(text), "used_chars": len(body)}


def extract_json(raw: str):
    s = str(raw).strip()
    s = re.sub(r"^```json\s*", "", s, flags=re.I)
    s = re.sub(r"^```\s*", "", s)
    s = re.sub(r"```\s*$", "", s)
    a, b = s.find("{"), s.rfind("}")
    if a >= 0 and b > a:
        s = s[a : b + 1]
    return json.loads(s)


# ─────────────────────────── FastAPI 应用 ───────────────────────────
app = FastAPI(title="明鉴 · AI 评审打分 API", version="1.8.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],           # 部署后建议收紧为你的前端域名
    allow_methods=["*"],
    allow_headers=["*"],
)


class ReviewRequest(BaseModel):
    track: str = Field(..., description="赛道 key，如 gaojiao_chuangyi")
    text: str = Field(..., min_length=50, description="从计划书提取的正文（至少 50 字）")
    material: str = Field("bp", description="bp | ppt")


@app.get("/api/health")
def health():
    ok = bool(DEEPSEEK_API_KEY)
    return {"status": "ok", "engine": "mingjian-v1.8.0", "deepseek_key_configured": ok}


@app.get("/api/tracks")
def tracks():
    return {"tracks": [{"key": k, "name": REVIEWER_DATA["tracks"][k]["name"]} for k in TRACK_KEYS]}


@app.post("/api/review")
async def review(req: ReviewRequest):
    if not DEEPSEEK_API_KEY:
        raise HTTPException(status_code=500, detail="服务端未配置 DEEPSEEK_API_KEY 环境变量")

    if req.track not in REVIEWER_DATA["tracks"]:
        raise HTTPException(status_code=400, detail=f"未知赛道: {req.track}")

    text = req.text.strip()
    if len(text) < 50:
        raise HTTPException(status_code=400, detail="计划书内容不足：请上传 PDF / Word 文件提取文字（至少 50 字）")

    # 超长文档（超过该赛道正文预算且 > 2.5 万字）→ 分段提取 + 多轮评审
    segmented = False
    segments = 1
    if len(text) > body_budget_for(req.track) and len(text) > 25000:
        log.info("超长文档，进入分段评审: 长度=%d", len(text))
        text, segments = await segment_review_text(text)
        segmented = True

    user_prompt, trunc_info = build_prompt(req.track, text)
    log.info("评审请求: track=%s material=%s 正文长度=%d prompt长度=%d 截断=%s 分段=%s(%d段)",
             req.track, req.material, len(text), len(user_prompt), trunc_info["truncated"], segmented, segments)

    payload = {
        "model": DEEPSEEK_MODEL,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": 0.2,
        "max_tokens": 4000,
        "response_format": {"type": "json_object"},
        "stream": False,
    }

    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            resp = await client.post(
                DEEPSEEK_URL,
                headers={
                    "Content-Type": "application/json",
                    "Authorization": "Bearer " + DEEPSEEK_API_KEY,
                },
                json=payload,
            )
    except httpx.HTTPError as e:
        log.error("DeepSeek 请求失败: %s", e)
        raise HTTPException(status_code=502, detail="上游模型服务请求失败: " + str(e))

    if resp.status_code != 200:
        log.error("DeepSeek 返回 %s: %s", resp.status_code, resp.text[:300])
        raise HTTPException(status_code=502, detail=f"上游模型服务返回 HTTP {resp.status_code}")

    data = resp.json()
    content = (
        data.get("choices", [{}])[0]
        .get("message", {})
        .get("content", "")
    )
    if not content:
        raise HTTPException(status_code=502, detail="模型返回内容为空")

    try:
        result = extract_json(content)
    except json.JSONDecodeError as e:
        log.error("JSON 解析失败: %s\n原始内容: %s", e, content[:500])
        raise HTTPException(status_code=502, detail="模型输出解析失败，请重试")

    # 轻量校验：total 必须是数字
    if "total" not in result:
        result["total"] = 0
    try:
        result["total"] = round(float(result["total"]), 1)
    except (TypeError, ValueError):
        result["total"] = 0

    return {
        "ok": True,
        "track": req.track,
        "material": req.material,
        "truncated": trunc_info["truncated"],
        "input_chars": trunc_info["input_chars"],
        "used_chars": trunc_info["used_chars"],
        "segmented": segmented,
        "segments": segments,
        "result": result,
    }


# ─────────────────────────── Word 报告导出（明鉴专业商业报告格式） ───────────────────────────
class ExportRequest(BaseModel):
    title: str = Field("创新作品", description="作品名称（用于报告封面与文件名）")
    track: str = Field("", description="赛道展示名，如 高教主赛道 · 创意组")
    material: str = Field("bp", description="bp | ppt")
    result: dict = Field(..., description="/api/review 返回的 result 对象")


def _set_run_font(run, name="微软雅黑", size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = _to_rgb(color)


def _add_para(doc, text, size=12, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    _set_run_font(p.add_run(text), size=size, bold=bold, color=color)
    if align is not None:
        p.alignment = align
    return p


def _add_table(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ""
        _set_run_font(cell.paragraphs[0].add_run(h), size=10.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            _set_run_font(cells[i].paragraphs[0].add_run(str(val)), size=10.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def build_report_docx(title, track_name, material, r) -> bytes:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()

    # 默认正文样式
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ── 封面信息页 ──
    _add_para(doc, "", size=12)
    _add_para(doc, "创新竞赛评审报告", size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "明鉴 · 创新竞赛评审专家", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "", size=12)
    _add_para(doc, f"作品名称：{title}", size=13)
    _add_para(doc, f"评审模式：官方竞赛规则（模式 A）", size=13)
    _add_para(doc, f"赛道：{track_name or '—'}", size=13)
    _add_para(doc, f"材料类型：{'PPT' if material == 'ppt' else 'BP（商业计划书）'}", size=13)
    _add_para(doc, f"总分：{r.get('total', '—')} / 100（得分率 {r.get('total', 0)}%）", size=13)
    _add_para(doc, f"档位：{r.get('level', '—')}（{r.get('level_range', '—')}）", size=13)
    if r.get("work_score") is not None:
        _add_para(doc, f"作品水平分：{r['work_score']} / 100（双轨评分）", size=13)
    _add_para(doc, f"评审日期：{date.today().isoformat()}", size=13)
    _add_para(doc, "报告版本：V1.0", size=13)
    doc.add_page_break()

    # ── 执行摘要 ──
    doc.add_heading("执行摘要", level=1)
    _add_para(doc, f"总体结论：{r.get('summary', '—')}", size=12, bold=True)
    if r.get("max_leverage"):
        _add_para(doc, f"最大改造杠杆：{r['max_leverage']}", size=12)

    # ── 一、评分总览 ──
    doc.add_heading("一、评分总览", level=1)
    rows = [[d.get("name", ""), d.get("full", ""), d.get("score", ""), d.get("diagnosis", "")] for d in r.get("dims", [])]
    rows.append(["总分", "100", r.get("total", ""), f"{r.get('level', '')}（{r.get('level_range', '')}）"])
    _add_table(doc, ["维度", "满分", "得分", "一句话诊断"], rows, widths=[3.2, 1.8, 1.8, 10.2])

    # ── 二、分项评分（逐维度） ──
    doc.add_heading("二、分项评分（逐维度）", level=1)
    for d in r.get("dims", []):
        doc.add_heading(f"{d.get('name', '')}（{d.get('score', '—')} / {d.get('full', '—')} 分）", level=2)
        if d.get("diagnosis"):
            _add_para(doc, f"关键判断：{d['diagnosis']}", size=11)
        subs = [
            [s.get("name", ""), f"{s.get('score', '—')} / {s.get('full', '—')}", s.get("loss", "无")]
            for s in d.get("subdims", [])
        ]
        if subs:
            _add_table(doc, ["子维度", "得分", "关键失分点"], subs, widths=[4.0, 2.5, 10.5])

    # ── 三、核心亮点 ──
    doc.add_heading("三、核心亮点", level=1)
    for i, h in enumerate(r.get("highlights", []) or ["（无）"], 1):
        _add_para(doc, f"{i}. {h}", size=12)

    # ── 四、关键不足 ──
    doc.add_heading("四、关键不足", level=1)
    for i, w in enumerate(r.get("weaknesses", []) or ["（无）"], 1):
        _add_para(doc, f"{i}. {w}", size=12)

    # ── 五、改进建议 ──
    doc.add_heading("五、改进建议（按优先级）", level=1)
    for s in r.get("suggestions", []):
        p = doc.add_paragraph()
        _set_run_font(p.add_run(f"优先级 {s.get('priority', '—')}："), size=12, bold=True)
        _set_run_font(p.add_run(str(s.get("text", ""))), size=12)

    # ── 六、获奖潜力 ──
    doc.add_heading("六、获奖潜力判断", level=1)
    _add_para(doc, str(r.get("potential", "—")), size=12)

    _add_para(doc, "", size=12)
    _add_para(doc, "本报告由「明鉴」AI 评审引擎自动生成，评分仅供参考，最终以大赛官方评审为准。", size=10, color=0x888888)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_sheet_docx(title, track_name, material, r) -> bytes:
    """评分明细表（独立文档）：逐小项展示得分点证据与扣分点"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ── 封面信息页 ──
    _add_para(doc, "", size=12)
    _add_para(doc, "评分明细表", size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "明鉴 · 创新竞赛评审专家", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "", size=12)
    _add_para(doc, f"作品名称：{title}", size=13)
    _add_para(doc, f"赛道：{track_name or '—'}", size=13)
    _add_para(doc, f"材料类型：{'PPT' if material == 'ppt' else 'BP（商业计划书）'}", size=13)
    _add_para(doc, f"总分：{r.get('total', '—')} / 100（{r.get('level', '—')}，{r.get('level_range', '—')}）", size=13)
    if r.get("work_score") is not None:
        _add_para(doc, f"作品水平分：{r['work_score']} / 100（双轨评分）", size=13)
    _add_para(doc, f"评审日期：{date.today().isoformat()}", size=13)
    doc.add_page_break()

    # ── 逐维度明细表（5 列：小项|满分|得分|得分点|扣分点） ──
    for d in r.get("dims", []):
        doc.add_heading(f"{d.get('name', '')}（{d.get('score', '—')} / {d.get('full', '—')} 分）", level=1)
        if d.get("diagnosis"):
            _add_para(doc, f"维度小结：{d['diagnosis']}", size=11)
        subs = [
            [s.get("name", ""), s.get("full", "—"), s.get("score", "—"), s.get("gain", "—"), s.get("loss", "无")]
            for s in d.get("subdims", [])
        ]
        if subs:
            _add_table(doc, ["小项", "满分", "得分", "得分点（证据）", "扣分点"], subs, widths=[2.6, 1.4, 1.4, 6.3, 5.3])

    # ── 总分汇总表 ──
    doc.add_heading("总分汇总", level=1)
    total_rows = [[d.get("name", ""), d.get("full", ""), d.get("score", "")] for d in r.get("dims", [])]
    total_rows.append(["总分", "100", r.get("total", "")])
    _add_table(doc, ["维度", "满分", "得分"], total_rows, widths=[4.0, 2.5, 2.5])

    # ── 档位判定 ──
    doc.add_heading("档位判定", level=1)
    _add_para(doc, f"当前档位：{r.get('level', '—')}（{r.get('level_range', '—')}）", size=12)
    if r.get("potential"):
        _add_para(doc, f"获奖潜力：{r['potential']}", size=12)

    _add_para(doc, "", size=12)
    _add_para(doc, "本表由「明鉴」AI 评审引擎自动生成，逐小项得分点（证据）与扣分点仅供参考，最终以大赛官方评审为准。", size=10, color=0x888888)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_response(content, title, suffix):
    safe = re.sub(r'[\\/:*?"<>|]', "_", title)[:40] or "创新作品"
    filename = f"{safe}_{suffix}_{date.today().strftime('%Y%m%d')}.docx"
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename*=UTF-8''" + quote(filename)},
    )


@app.post("/api/export_docx")
def export_docx(req: ExportRequest):
    if not HAS_DOCX:
        raise HTTPException(status_code=500, detail="服务端未安装 python-docx 依赖")
    try:
        content = build_report_docx(req.title, req.track, req.material, req.result)
    except Exception as e:
        log.error("docx 生成失败: %s", e)
        raise HTTPException(status_code=500, detail="Word 文档生成失败: " + str(e))
    return _docx_response(content, req.title, "创新竞赛评审报告")


@app.post("/api/export_sheet")
def export_sheet(req: ExportRequest):
    """评分明细表（独立文档）"""
    if not HAS_DOCX:
        raise HTTPException(status_code=500, detail="服务端未安装 python-docx 依赖")
    try:
        content = build_sheet_docx(req.title, req.track, req.material, req.result)
    except Exception as e:
        log.error("评分明细表生成失败: %s", e)
        raise HTTPException(status_code=500, detail="评分明细表生成失败: " + str(e))
    return _docx_response(content, req.title, "评分明细表")


if __name__ == "__main__":
    import uvicorn

    port = int(os.environ.get("PORT", "8000"))
    uvicorn.run(app, host="0.0.0.0", port=port)
